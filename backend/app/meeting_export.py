"""
Meeting export to Markdown / DOCX.

We deliberately skip PDF in v1: doing it well in mixed Chinese+English
needs an embedded CJK font (~10MB) and a heavy library (weasyprint
needs cairo/pango/gdk-pixbuf system deps; reportlab needs a registered
TTF). DOCX gives users a fully-formatted file they can open in Word
and "Save as PDF" themselves — same workflow at zero install cost.

Output in both formats:
  # <Meeting title>
  日期 / 参会人 / 状态

  ## 会议纪要
  <summary markdown verbatim>

  ## 实录
  [mm:ss] <speaker name>: <text>
  ...

  ## AI 专家发言
  【<agent name>】<text>
"""

from __future__ import annotations

import io
import logging
import uuid
from typing import Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from .models import (
    Agent,
    Meeting,
    MeetingAgentMessage,
    MeetingAttendee,
    MeetingTranscript,
    User,
)

logger = logging.getLogger(__name__)


def _fmt_ms(ms: Optional[int]) -> str:
    if ms is None:
        return "  ?  "
    s = ms / 1000.0
    return f"{int(s // 60):02d}:{int(s % 60):02d}"


async def _gather(meeting_id: uuid.UUID, db: AsyncSession) -> dict:
    meeting = (
        await db.execute(select(Meeting).where(Meeting.id == meeting_id))
    ).scalar_one_or_none()
    if not meeting:
        return {}

    transcripts = (
        await db.execute(
            select(MeetingTranscript)
            .where(MeetingTranscript.meeting_id == meeting_id)
            .order_by(MeetingTranscript.id)
        )
    ).scalars().all()

    user_ids = {t.speaker_user_id for t in transcripts if t.speaker_user_id}
    name_by_user: dict[uuid.UUID, str] = {}
    if user_ids:
        users = (
            await db.execute(select(User).where(User.id.in_(user_ids)))
        ).scalars().all()
        name_by_user = {u.id: u.name for u in users}

    attendee_rows = (
        await db.execute(
            select(MeetingAttendee).where(MeetingAttendee.meeting_id == meeting_id)
        )
    ).scalars().all()
    attendee_user_ids = [a.user_id for a in attendee_rows if a.user_id]
    if attendee_user_ids:
        users = (
            await db.execute(select(User).where(User.id.in_(attendee_user_ids)))
        ).scalars().all()
        attendee_names = [u.name for u in users]
    else:
        attendee_names = []

    agent_msgs = (
        await db.execute(
            select(MeetingAgentMessage)
            .where(MeetingAgentMessage.meeting_id == meeting_id)
            .order_by(MeetingAgentMessage.id)
        )
    ).scalars().all()
    agent_ids = {m.agent_id for m in agent_msgs}
    name_by_agent: dict[uuid.UUID, str] = {}
    if agent_ids:
        agents = (
            await db.execute(select(Agent).where(Agent.id.in_(agent_ids)))
        ).scalars().all()
        name_by_agent = {a.id: a.name for a in agents}

    return {
        "meeting": meeting,
        "transcripts": transcripts,
        "name_by_user": name_by_user,
        "attendee_names": attendee_names,
        "agent_msgs": agent_msgs,
        "name_by_agent": name_by_agent,
    }


# ---- Markdown ----------------------------------------------------------------

async def export_markdown(meeting_id: uuid.UUID, db: AsyncSession) -> Optional[str]:
    bundle = await _gather(meeting_id, db)
    if not bundle:
        return None
    m = bundle["meeting"]

    parts: list[str] = []
    parts.append(f"# {m.title}\n")

    meta_lines = []
    if m.started_at:
        meta_lines.append(f"- 时间：{m.started_at.strftime('%Y-%m-%d %H:%M')}")
    if bundle["attendee_names"]:
        meta_lines.append(f"- 参会人：{'、'.join(bundle['attendee_names'])}")
    meta_lines.append(f"- 状态：{m.status}")
    parts.append("\n".join(meta_lines) + "\n")

    if m.summary_md and not m.summary_md.startswith("<!--"):
        parts.append("## 会议纪要\n")
        parts.append(m.summary_md.strip() + "\n")

    if bundle["transcripts"]:
        parts.append("## 实录\n")
        for t in bundle["transcripts"]:
            speaker = (
                bundle["name_by_user"].get(t.speaker_user_id)
                if t.speaker_user_id
                else "未识别"
            )
            parts.append(f"- `[{_fmt_ms(t.start_ms)}]` **{speaker}**：{t.text.strip()}")
        parts.append("")

    if bundle["agent_msgs"]:
        parts.append("## AI 专家发言\n")
        for am in bundle["agent_msgs"]:
            agent_name = bundle["name_by_agent"].get(am.agent_id, "AI 专家")
            parts.append(f"### 🤖 {agent_name}\n")
            parts.append(am.text.strip() + "\n")

    return "\n".join(parts).strip() + "\n"


# ---- DOCX --------------------------------------------------------------------

async def export_docx(meeting_id: uuid.UUID, db: AsyncSession) -> Optional[bytes]:
    bundle = await _gather(meeting_id, db)
    if not bundle:
        return None
    m = bundle["meeting"]

    from docx import Document
    from docx.shared import Pt

    doc = Document()

    title = doc.add_heading(m.title, level=0)

    meta_lines = []
    if m.started_at:
        meta_lines.append(f"时间：{m.started_at.strftime('%Y-%m-%d %H:%M')}")
    if bundle["attendee_names"]:
        meta_lines.append(f"参会人：{'、'.join(bundle['attendee_names'])}")
    meta_lines.append(f"状态：{m.status}")
    meta_para = doc.add_paragraph(" · ".join(meta_lines))
    for run in meta_para.runs:
        run.font.size = Pt(9)

    # Summary section
    if m.summary_md and not m.summary_md.startswith("<!--"):
        doc.add_heading("会议纪要", level=1)
        # Render markdown line-by-line. We honor `## H2`/`### H3`/`- bullet`/
        # plain text. No external markdown parser dep.
        for line in m.summary_md.splitlines():
            stripped = line.rstrip()
            if not stripped:
                continue
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:].strip(), level=3)
            elif stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
                # GFM task list — render as bullet with checkbox prefix
                check = "☐ " if stripped.startswith("- [ ] ") else "☑ "
                doc.add_paragraph(check + stripped[6:].strip(), style="List Bullet")
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    # Transcript section
    if bundle["transcripts"]:
        doc.add_heading("实录", level=1)
        for t in bundle["transcripts"]:
            speaker = (
                bundle["name_by_user"].get(t.speaker_user_id)
                if t.speaker_user_id
                else "未识别"
            )
            p = doc.add_paragraph()
            ts_run = p.add_run(f"[{_fmt_ms(t.start_ms)}] ")
            ts_run.font.size = Pt(9)
            sp_run = p.add_run(f"{speaker}：")
            sp_run.bold = True
            p.add_run(t.text.strip())

    # Agent contributions
    if bundle["agent_msgs"]:
        doc.add_heading("AI 专家发言", level=1)
        for am in bundle["agent_msgs"]:
            agent_name = bundle["name_by_agent"].get(am.agent_id, "AI 专家")
            doc.add_heading(f"🤖 {agent_name}", level=2)
            doc.add_paragraph(am.text.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
