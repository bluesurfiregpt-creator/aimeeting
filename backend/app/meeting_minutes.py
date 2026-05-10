"""
v25-5 — 单场会议纪要 docx 导出.

智慧住建文档隐含需求(政务公文):客户开完会需要把纪要 docx 带走 + 留档.

内容(按顺序):
  1. 标题:会议纪要 — {会议标题}
  2. 元数据:开始 / 结束时间 / 时长 / 状态 / 参会人数
  3. 议程 agenda(若有)
  4. 摘要 summary_md(若有)
  5. 参会人员 attendees(用户 + AI 专家分两栏)
  6. 转写正文 transcript(分话人,带时间戳;长会话仅取关键发言或全部,看用户偏好)
  7. AI 发言 agent messages(按时间顺序穿插或单独成节)
  8. 待办事项 action items
  9. 页脚 — 系统名 + 生成时间
"""

from __future__ import annotations

import io
import logging
import uuid
from datetime import datetime, timezone
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from .models import (
    Agent,
    Meeting,
    MeetingActionItem,
    MeetingAgentMessage,
    MeetingAttendee,
    MeetingTranscript,
    User,
)

logger = logging.getLogger(__name__)


def _fmt_dt(dt: Optional[datetime]) -> str:
    """格式化时间为 北京时区 yyyy-mm-dd HH:MM."""
    if not dt:
        return "—"
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    # 转 北京时间(UTC+8)显示 — 政务公文场景下读者预期本地时间
    from datetime import timedelta
    beijing = timezone(timedelta(hours=8))
    return dt.astimezone(beijing).strftime("%Y-%m-%d %H:%M")


def _fmt_duration(start: Optional[datetime], end: Optional[datetime]) -> str:
    if not start or not end:
        return "—"
    if start.tzinfo is None:
        start = start.replace(tzinfo=timezone.utc)
    if end.tzinfo is None:
        end = end.replace(tzinfo=timezone.utc)
    delta = end - start
    total_min = int(delta.total_seconds() // 60)
    if total_min < 60:
        return f"{total_min} 分钟"
    return f"{total_min // 60} 小时 {total_min % 60} 分钟"


def _ms_to_clock(ms: Optional[int]) -> str:
    """0..3600000 → 'mm:ss'."""
    if ms is None:
        return ""
    sec = max(0, ms // 1000)
    return f"{sec // 60:02d}:{sec % 60:02d}"


async def build_minutes_docx(
    session: AsyncSession, meeting_id: uuid.UUID
) -> tuple[bytes, str]:
    """
    生成会议纪要 docx,返回 (bytes, suggested_filename_without_ext_中文).

    Filename 含会议标题 + 日期(中文,UTF-8;router 层做 RFC 5987 编码).
    """
    meeting = (
        await session.execute(select(Meeting).where(Meeting.id == meeting_id))
    ).scalar_one_or_none()
    if not meeting:
        raise ValueError(f"meeting {meeting_id} not found")

    # 拉关联数据
    attendees = (
        await session.execute(
            select(MeetingAttendee).where(MeetingAttendee.meeting_id == meeting_id)
        )
    ).scalars().all()
    user_ids = [a.user_id for a in attendees if a.user_id]
    agent_ids = [a.agent_id for a in attendees if a.agent_id]
    users_by_id: dict[uuid.UUID, User] = {}
    if user_ids:
        users = (
            await session.execute(select(User).where(User.id.in_(user_ids)))
        ).scalars().all()
        users_by_id = {u.id: u for u in users}
    agents_by_id: dict[uuid.UUID, Agent] = {}
    if agent_ids:
        agents = (
            await session.execute(select(Agent).where(Agent.id.in_(agent_ids)))
        ).scalars().all()
        agents_by_id = {a.id: a for a in agents}

    transcripts = (
        await session.execute(
            select(MeetingTranscript)
            .where(
                MeetingTranscript.meeting_id == meeting_id,
                MeetingTranscript.is_final.is_(True),
            )
            .order_by(MeetingTranscript.id.asc())
        )
    ).scalars().all()
    # transcript 用户名缓存
    transcript_user_ids = {t.speaker_user_id for t in transcripts if t.speaker_user_id}
    missing_uids = transcript_user_ids - set(users_by_id.keys())
    if missing_uids:
        more = (
            await session.execute(select(User).where(User.id.in_(missing_uids)))
        ).scalars().all()
        for u in more:
            users_by_id[u.id] = u

    agent_msgs = (
        await session.execute(
            select(MeetingAgentMessage)
            .where(MeetingAgentMessage.meeting_id == meeting_id)
            .order_by(MeetingAgentMessage.id.asc())
        )
    ).scalars().all()
    msg_agent_ids = {m.agent_id for m in agent_msgs}
    for aid in msg_agent_ids - set(agents_by_id.keys()):
        ag = (
            await session.execute(select(Agent).where(Agent.id == aid))
        ).scalar_one_or_none()
        if ag:
            agents_by_id[ag.id] = ag

    actions = (
        await session.execute(
            select(MeetingActionItem)
            .where(MeetingActionItem.meeting_id == meeting_id)
            .order_by(MeetingActionItem.created_at.asc())
        )
    ).scalars().all()

    # ---- 构造 docx --------------------------------------------------------
    doc = Document()

    # 全局字体(中文常用宋体;英文 Calibri)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    # 标题
    title = doc.add_heading(level=0)
    title_run = title.add_run("会议纪要")
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(meeting.title or "(未命名会议)")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # 元数据表
    doc.add_heading("一、会议信息", level=1)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Light List Accent 1"
    duration = _fmt_duration(meeting.started_at, meeting.ended_at)
    rows_data = [
        ("开始时间", _fmt_dt(meeting.started_at)),
        ("结束时间", _fmt_dt(meeting.ended_at)),
        ("时长", duration),
        ("状态", {"scheduled": "计划中", "ongoing": "进行中", "finished": "已结束", "processed": "已处理"}.get(meeting.status, meeting.status)),
    ]
    for i, (k, v) in enumerate(rows_data):
        cells = info_table.rows[i].cells
        cells[0].text = k
        cells[1].text = v

    # 议程
    if meeting.agenda:
        doc.add_heading("二、议程", level=1)
        for idx, item in enumerate(meeting.agenda, 1):
            t = item.get("title") if isinstance(item, dict) else str(item)
            tb = item.get("time_budget_min") if isinstance(item, dict) else None
            note = item.get("note") if isinstance(item, dict) else None
            line = f"{idx}. {t}"
            if tb:
                line += f"({tb} 分钟)"
            p = doc.add_paragraph(line)
            if note:
                p.add_run(f"\n   备注:{note}").italic = True

    # 参会人员
    next_section_index = 3 if meeting.agenda else 2
    doc.add_heading(f"{_chinese_num(next_section_index)}、参会人员", level=1)
    next_section_index += 1
    user_names = sorted([
        users_by_id[a.user_id].name
        for a in attendees
        if a.user_id and a.user_id in users_by_id
    ])
    agent_names = sorted([
        agents_by_id[a.agent_id].name
        for a in attendees
        if a.agent_id and a.agent_id in agents_by_id
    ])
    if user_names:
        doc.add_paragraph(f"参会人员({len(user_names)}):" + "、".join(user_names))
    else:
        doc.add_paragraph("(无登记参会人员)")
    if agent_names:
        doc.add_paragraph(f"参会 AI 专家({len(agent_names)}):" + "、".join(agent_names))

    # 摘要
    if meeting.summary_md:
        doc.add_heading(f"{_chinese_num(next_section_index)}、会议摘要", level=1)
        next_section_index += 1
        # 简化的 markdown 渲染:## → heading2, **bold** → bold run
        for line in meeting.summary_md.split("\n"):
            line = line.rstrip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif line.startswith(("1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. ", "8. ", "9. ")):
                doc.add_paragraph(line[3:].strip(), style="List Number")
            else:
                _add_paragraph_with_inline_format(doc, line)

    # 转写正文(限 200 行)
    if transcripts:
        doc.add_heading(f"{_chinese_num(next_section_index)}、会议记录(发言摘要)", level=1)
        next_section_index += 1
        MAX_LINES = 200
        shown = transcripts[:MAX_LINES]
        for t in shown:
            spk = (
                users_by_id[t.speaker_user_id].name
                if t.speaker_user_id and t.speaker_user_id in users_by_id
                else "未识别"
            )
            ts = _ms_to_clock(t.start_ms)
            p = doc.add_paragraph()
            r1 = p.add_run(f"[{ts}] {spk}:")
            r1.bold = True
            r1.font.color.rgb = RGBColor(0x55, 0x55, 0x88)
            p.add_run(f" {t.text}")
        if len(transcripts) > MAX_LINES:
            p = doc.add_paragraph()
            p.add_run(
                f"…(共 {len(transcripts)} 条,本纪要仅展示前 {MAX_LINES} 条;"
                f"完整记录见系统会议详情页)"
            ).italic = True

    # AI 发言
    if agent_msgs:
        doc.add_heading(f"{_chinese_num(next_section_index)}、AI 专家发言", level=1)
        next_section_index += 1
        for m in agent_msgs:
            ag_name = agents_by_id[m.agent_id].name if m.agent_id in agents_by_id else "AI 专家"
            p = doc.add_paragraph()
            r1 = p.add_run(f"【{ag_name}】 ")
            r1.bold = True
            r1.font.color.rgb = RGBColor(0x6B, 0x46, 0xC1)  # violet
            p.add_run(m.text or "")
            if m.citations:
                cite_p = doc.add_paragraph()
                cite_p.paragraph_format.left_indent = Pt(20)
                cr = cite_p.add_run("引用:")
                cr.italic = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                for c in m.citations[:3]:
                    fname = c.get("document_filename", "(未知文件)") if isinstance(c, dict) else "?"
                    cr2 = cite_p.add_run(f" 《{fname}》")
                    cr2.italic = True
                    cr2.font.size = Pt(9)
                    cr2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 待办事项
    if actions:
        doc.add_heading(f"{_chinese_num(next_section_index)}、待办事项", level=1)
        next_section_index += 1
        for idx, ai in enumerate(actions, 1):
            assignee = ""
            if ai.assignee_user_id and ai.assignee_user_id in users_by_id:
                assignee = users_by_id[ai.assignee_user_id].name
            elif ai.assignee_name_hint:
                assignee = ai.assignee_name_hint
            line = f"{idx}. {ai.content}"
            if assignee:
                line += f" — 负责人:{assignee}"
            if ai.due_at:
                line += f" — 截止:{_fmt_dt(ai.due_at)}"
            line += f" — 状态:{ {'open': '待办', 'done': '已完成', 'cancelled': '已取消'}.get(ai.status, ai.status) }"
            doc.add_paragraph(line, style="List Number")

    # 页脚
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = foot.add_run(
        f"\n— 本纪要由智慧住建 AI 集群系统自动生成 · "
        f"导出时间 {datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M')}"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    fr.italic = True

    # 输出
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    # 文件名:会议纪要-{title}-{date}.docx
    date_part = (
        meeting.started_at.strftime("%Y%m%d")
        if meeting.started_at
        else datetime.now().strftime("%Y%m%d")
    )
    safe_title = (meeting.title or "未命名会议")[:40]
    # Replace path-illegal chars
    for ch in '/\\:*?"<>|':
        safe_title = safe_title.replace(ch, "_")
    filename = f"会议纪要-{safe_title}-{date_part}.docx"

    return buf.getvalue(), filename


_CHINESE_NUMS = "一二三四五六七八九十"


def _chinese_num(n: int) -> str:
    """1→一,2→二,...10→十.超过 10 直接返回数字."""
    if 1 <= n <= 10:
        return _CHINESE_NUMS[n - 1]
    return str(n)


def _add_paragraph_with_inline_format(doc, line: str) -> None:
    """简单 inline markdown:**bold** 段渲染成粗体."""
    p = doc.add_paragraph()
    parts = line.split("**")
    bold = False
    for part in parts:
        run = p.add_run(part)
        run.bold = bold
        bold = not bold
